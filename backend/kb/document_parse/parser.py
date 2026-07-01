"""文档解析工具"""
from __future__ import annotations

import os
import re
import shutil
import tempfile
from typing import List, Tuple, Optional
from docx import Document as DocxDocument

from common.logging import logger

from kb.document_parse.factory import ParserFactory


class DocumentParser:
    @staticmethod
    def parse_file(
        file_path: str,
        domain: Optional[str] = None,
        business: Optional[str] = None,
        process_images: bool = False,
        parser_id: Optional[str] = None,
    ):
        """
        知识库统一文档解析入口（仅 parse，不分块）。唯一实现：DeepDoc。
        """
        _ = process_images  # DeepDoc 路径暂不使用 markitdown VLM 预处理
        return ParserFactory.parse(
            file_path,
            domain=domain,
            business=business,
            parser_id=parser_id,
        )

    @staticmethod
    def convert_file_to_markdown(file_path: str) -> str:
        """聊天附件等场景：优先 DeepDoc，失败时保留 markitdown 轻量转换。"""
        ext = os.path.splitext(file_path)[1].lower()
        if ext in {".md", ".markdown", ".mdown"}:
            return DocumentParser._read_markdown_file(file_path)
        try:
            parsed = DocumentParser.parse_file(file_path)
            text = (parsed.clean_markdown or parsed.raw_markdown or "").strip()
            if text:
                return text
        except Exception as exc:
            logger.warning(f"[DocumentParser] DeepDoc 转 Markdown 失败，尝试 markitdown: {exc}")
        try:
            from markitdown import MarkItDown

            result = MarkItDown().convert(file_path)
            return result.text_content or ""
        except Exception as exc:
            logger.error(f"文档转 Markdown 失败 {file_path}: {exc}")
            return ""

    @staticmethod
    def _read_markdown_file(file_path: str) -> str:
        text: Optional[str] = None
        for enc in ("utf-8-sig", "utf-8", "gbk"):
            try:
                with open(file_path, "r", encoding=enc) as f:
                    text = f.read()
                break
            except UnicodeDecodeError:
                continue
        if text is None:
            with open(file_path, "rb") as f:
                text = f.read().decode("utf-8", errors="replace")
        return text or ""

    @staticmethod
    def enhance_docx_to_markdown(
        file_path: str,
        process_images: bool = False,
        *,
        mutate_original: bool = True,
    ) -> Tuple[str, str]:
        """
        将 docx 文件转换为 markdown，并清理目录区域。
        
        当 process_images=True 时，提取文档中的 data URI 图片，调用 VL 模型生成描述，
        将描述替换回 markdown 中，得到可检索的 clean_text。
        
        Returns:
            (raw_markdown, clean_markdown)
            raw_markdown: 保留 data URI 的原始 markdown（用于溯源）
            clean_markdown: 图片替换为描述后的 markdown（用于 RAG 检索）
        """
        work_path = file_path
        temp_path: Optional[str] = None
        if not mutate_original:
            fd, temp_path = tempfile.mkstemp(suffix=os.path.splitext(file_path)[1] or ".docx")
            os.close(fd)
            shutil.copy2(file_path, temp_path)
            work_path = temp_path

        try:
            DocumentParser.wrap_code_blocks_in_docx(work_path)
            from markitdown import MarkItDown

            if process_images:
                md = MarkItDown()
                result = md.convert(work_path, keep_data_uris=True)
            else:
                md = MarkItDown()
                result = md.convert(work_path)

            raw_markdown = result.text_content
            raw_markdown = DocumentParser.clean_table_of_contents(raw_markdown)

            if not process_images:
                return raw_markdown, raw_markdown

            clean_markdown = DocumentParser._replace_images_with_descriptions(raw_markdown)
            return raw_markdown, clean_markdown
        finally:
            if temp_path and os.path.exists(temp_path):
                os.unlink(temp_path)

    @staticmethod
    def wrap_code_blocks_in_docx(file_path: str) -> None:
        doc = DocxDocument(file_path)
        in_code_block = False

        for i in range(len(doc.paragraphs) - 1, -1, -1):
            para = doc.paragraphs[i]
            is_terminal = para.style.name == 'Terminal Display'

            if is_terminal and not in_code_block:
                para.insert_paragraph_before("```")
                in_code_block = True
            elif not is_terminal and in_code_block:
                para.insert_paragraph_before("```")
                in_code_block = False

        if in_code_block:
            doc.paragraphs[0].insert_paragraph_before("```")

        doc.save(file_path)
        logger.info(f"docx 代码围栏处理完成: {file_path}")

    @staticmethod
    def _extract_data_uris(text: str) -> List[str]:
        """从 markdown 文本中提取所有 data URI"""
        pattern = r'data:image/[^;]+;base64,[A-Za-z0-9+/=]+'
        return re.findall(pattern, text)

    @staticmethod
    def _replace_images_with_descriptions(markdown_content: str) -> str:
        """
        将 markdown 中的 data URI 图片替换为 VLM 生成的文字描述。
        未配置 VLM 时静默跳过，保留原始图片内容。
        """
        from kb.embedding import is_vlm_configured

        if not is_vlm_configured():
            return markdown_content

        # 匹配 markdown 图片语法中的 data URI：![...](data:image/...;base64,...)
        img_pattern = re.compile(
            r'!\[([^\]]*)\]\((data:image/[^;]+;base64,[A-Za-z0-9+/=]+)\)'
        )

        def replace_one(match: re.Match) -> str:
            alt_text = match.group(1)
            data_uri = match.group(2)
            logger.info(f"发现 data URI 图片，alt='{alt_text}', 大小={len(data_uri)} 字节")
            try:
                description = DocumentParser._call_vl_model(data_uri)
                logger.info(f"图片描述生成成功，alt='{alt_text}'，描述长度={len(description)}")
                return f"\n\n[图片描述]: {description}\n\n"
            except BaseException as e:
                logger.warning(f"VLM 调用失败，使用占位文本: {type(e).__name__}: {e}")
                return f"\n\n[图片描述]: （图片内容，alt={alt_text}）\n\n"

        return img_pattern.sub(replace_one, markdown_content)

    @staticmethod
    def _call_vl_model(data_uri: str) -> str:
        """调用 VLM 对图片生成文字描述（mermaid 或自然语言）。"""
        import httpx
        from openai import OpenAI
        from markitdown import MarkItDown
        from config.env import ModelConfig
        from kb.embedding import is_vlm_configured

        if not is_vlm_configured():
            raise ValueError("VLM 未配置")

        api_key = ModelConfig.vlm_model_api_key.strip()
        logger.info(f"开始调用 VLM 模型，data_uri 大小={len(data_uri)} 字节")
        client = OpenAI(
            api_key=api_key,
            base_url=ModelConfig.vlm_model_base_url,
            http_client=httpx.Client(
                timeout=httpx.Timeout(connect=10, read=120, write=30, pool=10),
            ),
        )
        vl_model = ModelConfig.vlm_model_name
        logger.info(f"使用 VLM 模型: {vl_model}")
        md = MarkItDown(
            llm_client=client,
            llm_model=vl_model,
            llm_prompt="请用中文描述这张图片。如果是时序图或者其他适合用mermaid描述的图表，请使用mermaid描述其表达的技术逻辑。",
        )
        result = md.convert(data_uri)
        logger.info(f"VL 模型返回结果，长度={len(result.text_content or '')}")
        return result.text_content
    
    @staticmethod
    def clean_table_of_contents(markdown_content: str) -> str:
        """
        清理 Markdown 内容中的目录区域和目录链接
        
        策略：
        1. 直接删除所有包含 #_Toc 的目录链接行（更激进的清理）
        2. 检测连续的目录行并整体删除
        """
        lines = markdown_content.split('\n')
        
        # 第一步：定义 TOC 链接模式
        toc_pattern = r'\[.+?\]\(#_Toc\d+\)'
        
        # 第二步：统计并清理包含 TOC 链接的行
        cleaned_lines = []
        toc_link_count = 0
        
        for i, line in enumerate(lines):
            # 检测该行是否包含 TOC 链接
            if re.search(toc_pattern, line):
                toc_link_count += len(re.findall(toc_pattern, line))
                # 完全跳过包含 TOC 链接的行
                continue
            
            # 保留非 TOC 行
            cleaned_lines.append(line)
        
        if toc_link_count > 0:
            logger.info(f"共清理 {toc_link_count} 个目录链接")
        
        return '\n'.join(cleaned_lines)
